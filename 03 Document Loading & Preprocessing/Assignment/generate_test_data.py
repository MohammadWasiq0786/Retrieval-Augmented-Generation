import os
import json
import csv
from pathlib import Path
from docx import Document
from fpdf import FPDF

LOREM_PARAGRAPH = (
    "LangChain is a flexible framework designed for building context-aware applications powered by LLMs. "
    "It simplifies loading data from diverse sources, chunking text, generating embeddings, and retrieving results. "
    "With support for integrations like FAISS, OpenAI, HuggingFace, it enables developers to create RAG pipelines, "
    "chatbots, and AI agents seamlessly.\n"
)

def write_large_txt(file_path, paragraph, repeat=300):
    content = paragraph * repeat
    Path(file_path).write_text(content)
    print(f"[+] Generated large .txt file: {file_path}")

def write_large_csv(file_path, paragraph, repeat=300):
    with open(file_path, mode="w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["id", "description"])
        writer.writeheader()
        for i in range(1, repeat + 1):
            writer.writerow({"id": i, "description": paragraph.strip()})
    print(f"[+] Generated large .csv file: {file_path}")

def write_large_docx(file_path, paragraph, repeat=300):
    doc = Document()
    doc.add_heading("LangChain Deep Dive Report", 0)
    for _ in range(repeat):
        doc.add_paragraph(paragraph.strip())
    doc.save(file_path)
    print(f"[+] Generated large .docx file: {file_path}")

def write_large_pdf(file_path, paragraph, repeat=100):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for _ in range(repeat):
        pdf.multi_cell(0, 10, paragraph.strip())
    pdf.output(file_path)
    print(f"[+] Generated large .pdf file: {file_path}")

def write_notion_json(dir_path):
    os.makedirs(dir_path, exist_ok=True)
    pages = [
        {"title": f"Final Page {i}", "tags": ["final"], "content": LOREM_PARAGRAPH}
        for i in range(1, 6)
    ] + [
        {"title": f"Draft Page {i}", "tags": ["draft"], "content": LOREM_PARAGRAPH}
        for i in range(6, 9)
    ]
    for i, page in enumerate(pages):
        json_path = os.path.join(dir_path, f"page_{i+1}.json")
        with open(json_path, "w") as f:
            json.dump(page, f)
    print(f"[+] Generated Notion export files in: {dir_path}")

def create_test_data():
    os.makedirs("test_data", exist_ok=True)
    write_large_txt("test_data/large_sample.txt", LOREM_PARAGRAPH)
    write_large_csv("test_data/large_sample.csv", LOREM_PARAGRAPH)
    write_large_docx("test_data/large_sample.docx", LOREM_PARAGRAPH)
    write_large_pdf("test_data/large_sample.pdf", LOREM_PARAGRAPH)
    write_notion_json("test_data/notion_export")

if __name__ == "__main__":
    create_test_data()
